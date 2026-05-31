from __future__ import annotations

import sys
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from src.config import OUTPUT_DOC_DIR


def configure_styles(document: DocxDocument) -> None:
    for style_name in ["Normal", "List Bullet", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    document.styles["Title"].font.size = Pt(18)
    document.styles["Heading 1"].font.size = Pt(14)
    document.styles["Heading 2"].font.size = Pt(12)
    document.styles["Heading 3"].font.size = Pt(11)


def add_heading(document: DocxDocument, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: DocxDocument, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def build_document() -> DocxDocument:
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = document.add_heading("Resumo executivo - Desafio Conta Azul", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("Business Analytics - Produto SaaS")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)

    add_heading(document, "1. Principais achados do funil", level=1)
    add_paragraph(
        document,
        "A base possui 10.000 usuarios unicos entre 2025-06-01 e 2025-10-30. "
        "O funil geral tem 2.983 signups e 919 compras, com taxa de visita para signup "
        "de 29,83%, visita para compra de 9,19% e signup para compra de 30,81%.",
    )
    add_paragraph(
        document,
        "Referral e o canal de maior qualidade: 41,65% de visit to signup, 18,23% "
        "de visit to purchase e 43,77% de signup to purchase. Email tambem apresenta "
        "boa eficiencia, enquanto organic combina escala e conversao. Paid e social "
        "apresentam baixa conversao final.",
    )
    add_paragraph(
        document,
        "Mobile concentra 62,42% das visitas, mas converte menos que desktop. A taxa "
        "visit to purchase em mobile e 6,82%, contra 13,12% em desktop, indicando uma "
        "oportunidade relevante de melhoria de experiencia, cadastro, onboarding ou checkout.",
    )
    add_paragraph(
        document,
        "O NPS geral esconde diferencas importantes: compradores possuem NPS estimado "
        "de 33,33, enquanto nao compradores possuem NPS estimado de -29,46. Portanto, "
        "a experiencia dos usuarios que nao compram precisa ser investigada separadamente.",
    )

    add_heading(document, "2. Gargalos e hipoteses", level=1)
    add_paragraph(
        document,
        "O primeiro gargalo esta antes do signup: 70,17% dos visitantes nao criam conta. "
        "As hipoteses sao proposta de valor pouco clara, paginas de entrada pouco aderentes, "
        "friccao no cadastro ou baixa qualificacao de trafego em alguns canais.",
    )
    add_paragraph(
        document,
        "O segundo gargalo esta entre signup e purchase: 69,19% dos cadastrados nao assinam. "
        "As hipoteses sao onboarding insuficiente, demora para chegar ao momento de valor, "
        "pricing pouco claro, falta de prova social ou friccao no checkout.",
    )
    add_paragraph(
        document,
        "Paid tem alto volume, mas baixa eficiencia, sugerindo segmentacao, criativos ou "
        "landing pages desalinhadas. Mobile tambem merece prioridade por combinar alto volume "
        "com baixa conversao relativa.",
    )

    add_heading(document, "3. Recomendacoes de negocio", level=1)
    recommendations = [
        "Priorizar referral com programa de indicacao, incentivos controlados e landing page especifica.",
        "Otimizar paid antes de escalar investimento, medindo compra e qualidade em vez de apenas signup.",
        "Melhorar a jornada mobile, reduzindo friccao no cadastro, onboarding e checkout.",
        "Fortalecer organic com SEO, conteudo orientado ao ICP e paginas de conversao mais contextualizadas.",
        "Explorar email/lifecycle para nutrir usuarios pos-visita e recuperar signups sem compra.",
        "Investigar nao compradores com NPS baixo para identificar objeções, friccoes e lacunas de valor.",
    ]
    for item in recommendations:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, "Metodologia", level=1)
    add_paragraph(
        document,
        "A solucao foi desenvolvida localmente com Python, Pandas, DuckDB, Plotly e Streamlit. "
        "As metricas foram implementadas em SQL local com DuckDB e documentadas de forma agnostica, "
        "podendo ser reproduzidas em Looker, Power BI ou outra camada semantica equivalente.",
    )

    return document


def main() -> None:
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    output_path = OUTPUT_DOC_DIR / "resumo_executivo_conta_azul.docx"
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
