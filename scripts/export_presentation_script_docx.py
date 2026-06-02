from __future__ import annotations

import sys
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from src.config import OUTPUT_DOC_DIR

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


def configure_styles(document: "DocxDocument") -> None:
    for style_name in ["Normal", "List Bullet", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    document.styles["Title"].font.size = Pt(18)
    document.styles["Heading 1"].font.size = Pt(14)
    document.styles["Heading 2"].font.size = Pt(12)
    document.styles["Heading 3"].font.size = Pt(11)


def add_heading(document: "DocxDocument", text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: "DocxDocument", text: str, bold_label: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    if bold_label:
        label = paragraph.add_run(bold_label)
        label.bold = True
        label.font.name = "Calibri"
        label.font.size = Pt(10)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_bullet(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_quote(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def build_document() -> "DocxDocument":
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = document.add_heading("Roteiro de apresentacao - Desafio Conta Azul", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("Apoio para demonstracao do projeto de Business Analytics")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)

    add_heading(document, "1. Abertura", level=1)
    add_paragraph(
        document,
        "O objetivo do case era analisar o comportamento dos usuarios em um funil SaaS, "
        "identificar gargalos e oportunidades e transformar os achados em recomendacoes de negocio.",
    )
    add_paragraph(document, "Stack usada: ", bold_label="")
    for item in [
        "Pandas para leitura, profiling, tratamento e validacoes.",
        "DuckDB para views e consultas SQL locais.",
        "Plotly para graficos.",
        "Streamlit para dashboard executivo.",
        "Word para resumo formal do case.",
        "Notebook para demonstrar investigacao e raciocinio analitico.",
    ]:
        add_bullet(document, item)

    add_heading(document, "2. Ordem recomendada de demonstracao", level=1)
    for item in [
        "Mostrar rapidamente o PDF do desafio e reforcar os entregaveis pedidos.",
        "Abrir o notebook para mostrar a investigacao e o raciocinio analitico.",
        "Abrir o SQL 01_create_views.sql para mostrar as views DuckDB.",
        "Rodar ou mostrar o dashboard Streamlit.",
        "Mostrar o Word do resumo executivo como entregavel formal.",
        "Fechar na aba Resposta ao case, conectando a solucao ao resultado esperado.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3. Onde mostrar cada ponto no projeto", level=1)
    add_paragraph(document, "Notebook: ", bold_label="Investigacao e raciocinio - ")
    add_bullet(document, "Arquivo: notebooks/01_eda_funil_saas.ipynb")
    add_bullet(document, "Mostra leitura do CSV, profiling, validacoes, funil, segmentos, NPS e graficos Plotly.")

    add_paragraph(document, "SQL: ", bold_label="Consultas reproduziveis - ")
    add_bullet(document, "Arquivo principal: sql/01_create_views.sql")
    add_bullet(document, "Views: vw_funnel_overall, vw_funnel_by_channel, vw_funnel_by_device, vw_funnel_by_country, vw_funnel_by_month, vw_plan_mix, vw_nps_summary e vw_nps_by_channel.")

    add_paragraph(document, "Dashboard: ", bold_label="Leitura executiva - ")
    add_bullet(document, "Comando: py -3.13 -m streamlit run app.py")
    add_bullet(document, "Abas: Visao geral, Segmentos, NPS, Qualidade dos dados e Resposta ao case.")

    add_paragraph(document, "Word: ", bold_label="Resumo formal - ")
    add_bullet(document, "Arquivo: output/doc/resumo_executivo_conta_azul.docx")
    add_bullet(document, "Contem achados, gargalos/hipoteses e recomendacoes.")

    add_heading(document, "4. Como explicar o dashboard Streamlit", level=1)
    add_paragraph(
        document,
        "Pense no dashboard como a camada executiva da analise. O notebook mostra o raciocinio; "
        "o SQL mostra as metricas reproduziveis; o Streamlit mostra os resultados para tomada de decisao.",
    )
    add_quote(
        document,
        "O dashboard foi criado em Streamlit para transformar a analise em uma experiencia interativa. "
        "Ele permite filtrar por mes, canal, dispositivo e pais, e as metricas sao recalculadas a partir "
        "das consultas SQL no DuckDB.",
    )

    add_heading(document, "Filtros laterais", level=2)
    add_paragraph(document, "Na lateral esquerda tem:")
    for item in ["Mes da visita.", "Canal.", "Dispositivo.", "Pais."]:
        add_bullet(document, item)
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Os filtros permitem simular diferentes recortes do funil. Por exemplo, eu posso olhar apenas "
        "mobile, apenas paid ou um pais especifico. Isso ajuda a sair de uma media geral e investigar "
        "segmentos onde a conversao muda.",
    )

    add_heading(document, "Aba Visao geral", level=2)
    add_paragraph(document, "Essa aba responde: como esta o funil como um todo?")
    add_paragraph(document, "Ela mostra:")
    for item in [
        "Visitas.",
        "Signups.",
        "Compras.",
        "Signup -> compra.",
        "Respostas de NPS.",
        "NPS medio.",
        "Drop-off visita -> signup.",
        "Drop-off signup -> compra.",
        "Grafico de funil.",
        "Evolucao mensal das taxas.",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Na visao geral eu comeco pelo funil macro. A base tem 10.000 visitas, 2.983 signups e 919 compras. "
        "O principal gargalo esta antes do signup: 70,17% dos visitantes nao criam conta. O segundo gargalo "
        "esta depois do signup: 69,19% dos cadastrados nao compram.",
    )

    add_heading(document, "Aba Segmentos", level=2)
    add_paragraph(document, "Essa aba responde: onde estao as melhores e piores performances?")
    add_paragraph(document, "Ela mostra:")
    for item in [
        "Performance por canal.",
        "Tabela detalhada por canal.",
        "Conversao por dispositivo.",
        "Heatmap canal x dispositivo.",
        "Performance por pais.",
        "Mix por plano.",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Na aba de segmentos eu investigo onde a conversao melhora ou piora. Referral tem a melhor taxa "
        "visit to purchase, seguido por email e organic. Paid e social tem baixa eficiencia. Desktop "
        "converte melhor que mobile, mesmo mobile tendo grande volume.",
    )
    add_paragraph(document, "Sobre o grafico de canal:")
    add_quote(
        document,
        "Referral nao e o maior canal em volume, mas e o melhor em eficiencia. Organic traz volume e boa "
        "conversao. Paid traz volume, mas converte pouco; por isso eu nao recomendaria aumentar investimento "
        "antes de otimizar campanha, landing page e publico.",
    )
    add_paragraph(document, "Sobre o heatmap:")
    add_quote(
        document,
        "O heatmap cruza canal e dispositivo. Ele ajuda a identificar combinacoes especificas com melhor ou "
        "pior conversao, como referral no desktop performando muito bem e social/mobile performando pior.",
    )

    add_heading(document, "Aba NPS", level=2)
    add_paragraph(document, "Essa aba responde: como esta a experiencia dos usuarios?")
    add_paragraph(document, "Ela mostra:")
    for item in [
        "NPS de compradores elegiveis.",
        "Controle de respostas NPS nao elegiveis.",
        "Investigacao por canal, dispositivo, pais, signup e classe da nota.",
        "NPS por canal.",
        "Tabela detalhada.",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Aqui eu trato NPS como pesquisa pos-compra. Por isso, o indicador final considera apenas compradores "
        "elegiveis. As respostas associadas a usuarios sem compra aparecem como alerta para investigar regra de "
        "negocio, janela de atribuicao ou tracking.",
    )
    add_paragraph(document, "Sobre a investigacao das respostas nao elegiveis:")
    add_quote(
        document,
        "A investigacao mostra onde esses registros se concentram. Na base completa, aparecem mais em organic, "
        "paid, mobile e em usuarios que fizeram signup mas nao compraram. Isso ajuda a priorizar auditoria de "
        "tracking, janela de purchase e regra de disparo da pesquisa.",
    )
    add_paragraph(document, "Sobre NPS por canal:")
    add_quote(
        document,
        "Referral tambem aparece bem em NPS, o que reforca que e um canal de maior qualidade. Paid e social "
        "tem NPS mais baixo, entao podem estar trazendo usuarios menos qualificados ou com expectativa diferente.",
    )

    add_heading(document, "Aba Qualidade dos dados", level=2)
    add_paragraph(document, "Essa aba responde: posso confiar nos numeros?")
    add_paragraph(document, "Ela mostra:")
    for item in ["Profiling da base.", "Validacoes da base completa.", "Validacoes com filtros aplicados."]:
        add_bullet(document, item)
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Antes de apresentar recomendacoes, eu validei a qualidade dos dados. Verifiquei duplicidade de usuario, "
        "compra sem signup, compra sem plano, plano sem compra, NPS fora do intervalo e tempos inconsistentes. "
        "As regras criticas deram zero inconsistencias.",
    )
    add_paragraph(document, "Sobre `respostas_nps_nao_compradores = 465`:")
    add_quote(
        document,
        "Esse ponto nao entra no NPS final. Como NPS e pos-compra, essas respostas sem compra sao tratadas como "
        "ponto de investigacao: pode ser outro survey, compra fora da janela, atraso de tracking ou problema de "
        "instrumentacao.",
    )

    add_heading(document, "Aba Resposta ao case", level=2)
    add_paragraph(document, "Essa aba e a defesa direta do desafio.")
    add_paragraph(document, "Fala sugerida:")
    add_quote(
        document,
        "Essa aba conecta a solucao ao resultado esperado do case. Eu mostro como a analise entende o comportamento "
        "dos usuarios, identifica gargalos, aplica raciocinio analitico com validacoes e SQL, e transforma os achados "
        "em recomendacoes de negocio.",
    )
    add_paragraph(document, "Possivel resolucao adicionada ao case:")
    add_quote(
        document,
        "Para os registros NPS nao elegiveis, eu nao recomendo usar o numero como NPS de nao compradores. "
        "A resolucao analitica e manter o NPS apenas para compradores e abrir uma investigacao operacional: "
        "validar disparo da pesquisa, auditar tracking de compra, checar compras fora da janela e pesquisar abandono.",
    )
    add_paragraph(document, "Fechamento dessa aba:")
    add_quote(
        document,
        "O dashboard nao e so um conjunto de graficos. Ele organiza a investigacao do funil em uma narrativa: "
        "primeiro o tamanho do problema, depois onde ele acontece, depois a qualidade da experiencia e, por fim, "
        "quais acoes de negocio priorizar.",
    )

    add_heading(document, "Roteiro de demonstracao no Streamlit", level=2)
    for item in [
        "Comece em Visao geral: mostre visitas, signups, compras e drop-offs.",
        "Va para Segmentos: mostre referral, email, organic, paid e social.",
        "Ainda em Segmentos: mostre desktop vs mobile e o heatmap.",
        "Va para NPS: mostre NPS de compradores elegiveis, controle de elegibilidade, investigacao das respostas nao elegiveis e NPS por canal.",
        "Va para Qualidade dos dados: mostre que nao ha inconsistencias criticas.",
        "Termine em Resposta ao case: conecte tudo ao que o desafio pediu.",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "Frase boa para apresentacao:")
    add_quote(
        document,
        "Eu desenhei o dashboard para seguir a logica de decisao do negocio: primeiro entender o tamanho do funil, "
        "depois localizar onde a conversao muda, validar se os dados sao confiaveis e, por fim, traduzir os achados "
        "em recomendacoes praticas.",
    )

    add_heading(document, "5. Como responder ao resultado esperado", level=1)
    add_paragraph(document, "Entender comportamento dos usuarios: ", bold_label="")
    add_paragraph(
        document,
        "Eu medi o funil completo de visita para signup e de signup para compra. Depois segmentei "
        "o comportamento por canal, dispositivo, pais, mes, plano e NPS para entender onde a conversao "
        "melhora ou piora.",
    )

    add_paragraph(document, "Identificar gargalos e oportunidades: ", bold_label="")
    add_paragraph(
        document,
        "O principal gargalo esta antes do signup: 70,17% dos visitantes nao criam conta. O segundo "
        "gargalo esta entre signup e compra: 69,19% dos cadastrados nao compram. Tambem identifiquei "
        "oportunidades em mobile, paid e social, alem de investigar respostas NPS nao elegiveis.",
    )

    add_paragraph(document, "Aplicar raciocinio analitico e estatistico: ", bold_label="")
    add_paragraph(
        document,
        "Antes de concluir qualquer coisa, validei a base: duplicidades, compra sem signup, compra sem "
        "plano, NPS invalido e campos inconsistentes. Depois calculei as metricas com SQL no DuckDB e "
        "comparei taxas por segmento e calculei NPS apenas para compradores elegiveis.",
    )

    add_paragraph(document, "Transformar achados em recomendacoes: ", bold_label="")
    add_paragraph(
        document,
        "A partir dos achados, recomendei escalar referral, otimizar paid antes de aumentar investimento, "
        "melhorar mobile, fortalecer organic, usar email/lifecycle, investigar nao compradores por abandono "
        "e auditar respostas NPS nao elegiveis.",
    )

    add_heading(document, "6. Principais achados para falar", level=1)
    for item in [
        "Funil geral: 10.000 visitas, 2.983 signups e 919 compras.",
        "Visit to signup: 29,83%; visit to purchase: 9,19%; signup to purchase: 30,81%.",
        "Referral e o canal mais eficiente: 18,23% de visit to purchase.",
        "Email tem boa eficiencia, mas baixo volume.",
        "Organic combina escala e boa conversao.",
        "Paid tem alto volume, mas baixa eficiencia.",
        "Mobile tem mais visitas, mas converte menos que desktop.",
        "Compradores elegiveis tem NPS estimado de 33,33; respostas sem compra nao entram no indicador final.",
        "Respostas NPS nao elegiveis se concentram em organic, paid, mobile e usuarios com signup sem compra.",
    ]:
        add_bullet(document, item)

    add_heading(document, "7. Recomendacoes finais", level=1)
    for item in [
        "Escalar referral com programa de indicacao e controle de qualidade.",
        "Otimizar paid antes de aumentar budget.",
        "Melhorar UX, cadastro, onboarding e checkout mobile.",
        "Fortalecer organic com SEO e paginas orientadas ao ICP.",
        "Usar email/lifecycle para recuperar signups sem compra.",
        "Investigar nao compradores por pesquisa qualitativa e eventos de abandono.",
        "Auditar regra de disparo da pesquisa, tracking de purchase e janela de observacao para respostas NPS nao elegiveis.",
    ]:
        add_bullet(document, item)

    add_heading(document, "8. Frase de fechamento", level=1)
    add_paragraph(
        document,
        "Minha solucao nao ficou so em visualizacao. Eu estruturei uma trilha completa: qualidade dos dados, "
        "analise exploratoria, SQL reproduzivel, dashboard executivo e recomendacoes de negocio conectadas "
        "aos gargalos do funil.",
    )

    return document


def main() -> None:
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DOC_DIR / "roteiro_apresentacao_conta_azul.docx"
    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
