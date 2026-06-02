from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from src.config import OUTPUT_DOC_DIR


GUIDE_MARKDOWN = """
# Guia de apoio - Notebook `01_eda_funil_saas.ipynb`

Use este guia como uma cola de apresentacao. A ideia e explicar os resultados principais do notebook sem ler cada celula palavra por palavra.

Fala de abertura sugerida:

> Este notebook mostra a parte investigativa da analise. O dashboard traz a leitura executiva, mas aqui eu mostro como os dados foram carregados, validados, transformados e analisados antes de chegar nas recomendacoes.

## Pontos que podem gerar duvida

### Quando aparecer `WindowsPath(...)`

`WindowsPath(...)` e apenas o caminho local do CSV no Windows.

Fala sugerida:

> Esse retorno nao e erro. Ele apenas confirma o caminho do arquivo CSV usado na analise.

### Quando aparecer `(10000, 12)`

Esse resultado vem de `raw.shape` e representa:

- `10000`: linhas da base original.
- `12`: colunas originais do CSV.

Fala sugerida:

> Aqui eu confirmo que o CSV original tem 10.000 linhas e 12 colunas, exatamente a volumetria esperada para o desafio.

### Quando o profiling mostrar `colunas = 17`

O CSV original tem 12 colunas. Depois do tratamento, a base analitica ganha campos derivados: `visit_month`, `signup_date`, `purchase_date`, `funnel_stage` e `nps_class`.

Fala sugerida:

> A diferenca e intencional: 12 colunas e o formato original do CSV; 17 colunas e a base ja preparada para analise, com campos derivados de data, funil e NPS.

### Quando aparecer `NaN`

`NaN` significa valor ausente ou nao aplicavel. No funil, isso e esperado quando o usuario nao avancou para determinada etapa.

Fala sugerida:

> Esses nulos nao indicam erro por si so. Eles refletem a regra do funil: se o usuario nao avancou para uma etapa, alguns campos ficam naturalmente vazios.

## Celula 2 - Imports, caminhos e configuracao visual

Esta celula prepara o ambiente: importa bibliotecas, identifica a raiz do projeto, aponta para o CSV, configura as cores inspiradas na Conta Azul e define o renderer dos graficos Plotly.

Fala sugerida:

> Aqui preparo o ambiente para deixar a analise reproduzivel. O notebook encontra o CSV, carrega as bibliotecas e configura a identidade visual dos graficos.

## Celula 4 - Leitura e amostra inicial

Esta celula le o CSV com Pandas e mostra as primeiras linhas.

Fala sugerida:

> Primeiro eu leio o CSV e olho uma amostra da base para entender a estrutura original: usuario, data da visita, canal, dispositivo, pais, signup, compra, plano, tempos e NPS.

## Celula 5 - Dimensao da base original

Resultado:

- `linhas = 10000`
- `colunas = 12`

Fala sugerida:

> Aqui valido a volumetria original: 10.000 linhas e 12 colunas. Isso confirma que estou analisando a base completa do case.

## Celula 7 - Tratamento e campos derivados

Principais tratamentos: conversao de data, ajuste de tipos numericos, criacao de `visit_month`, datas estimadas, etapa final do funil e classificacao de NPS.

Fala sugerida:

> Aqui eu transformo o CSV original em uma base analitica. A ideia nao e alterar o dado bruto, mas criar campos que facilitem as analises de funil, tempo e NPS.

## Celula 9 - Profiling da base tratada

Resultado esperado:

- `linhas = 10000`
- `colunas = 17`
- `usuarios_unicos = 10000`
- `usuarios_duplicados = 0`
- `data_minima_visita = 2025-06-01`
- `data_maxima_visita = 2025-10-30`
- `canais = 5`
- `dispositivos = 2`
- `paises = 5`
- `planos = 3`

Fala sugerida:

> Essa tabela mostra que a base esta consistente: tenho 10.000 usuarios unicos, sem duplicidade, cobrindo visitas de junho a outubro de 2025.

## Celula 10 - Perfil das categorias

Codigo da celula:

```python
df[['channel', 'device', 'country', 'plan']].describe(include='all')
```

Como ler a tabela:

- `count`: quantidade de valores preenchidos.
- `unique`: quantidade de categorias diferentes.
- `top`: categoria mais frequente.
- `freq`: quantidade de vezes em que a categoria mais frequente aparece.

Resultado principal:

- `channel`: canal mais frequente e `organic`, com 4.271 usuarios.
- `device`: dispositivo mais frequente e `mobile`, com 6.242 usuarios.
- `country`: pais mais frequente e `BR`, com 6.986 usuarios.
- `plan`: plano mais frequente e `BASIC`, com 526 compradores.

Fala sugerida:

> Esta tabela mostra o mix da base. Organic e o canal mais comum, mobile e o dispositivo mais frequente, BR concentra a maior parte dos usuarios e BASIC e o plano mais comprado. O plano aparece em apenas 919 linhas porque so usuarios compradores possuem plano.

## Celula 12 - Validacoes de qualidade

Valida regras como compra sem signup, compra sem plano, plano sem compra, NPS fora do intervalo, tempos inconsistentes e usuarios duplicados.

Fala sugerida:

> Antes de tirar conclusoes, eu valido a qualidade dos dados. As regras criticas deram zero inconsistencias, entao as metricas do funil estao confiaveis.

## Celula 14 - DuckDB e views analiticas

A base tratada e registrada no DuckDB como `stg_funnel_users`, e o arquivo `sql/01_create_views.sql` cria as views analiticas.

Fala sugerida:

> Aqui eu trago SQL para a analise. Em vez de calcular tudo manualmente no notebook, registro a base no DuckDB e crio views reproduziveis para funil, segmentos, planos e NPS.

## Celula 16 - Funil geral

Principais resultados:

- visitas: 10.000;
- signups: 2.983;
- compras: 919;
- visit to signup: 29,83%;
- visit to purchase: 9,19%;
- signup to purchase: 30,81%.

Fala sugerida:

> Aqui esta o resumo do funil. De 10.000 visitantes, 2.983 criaram conta e 919 compraram. A conversao final sobre visitas foi de 9,19%.

## Celula 17 - Grafico de funil

Mostra visualmente a queda entre visitas, signups e compras.

Fala sugerida:

> O grafico deixa claro que existe uma perda grande antes do signup e outra perda relevante depois do signup.

## Celula 20 - Analise por canal

Leituras principais:

- `referral` tem melhor eficiencia;
- `email` tambem performa bem;
- `organic` combina volume e conversao;
- `paid` e `social` precisam de otimizacao.

Fala sugerida:

> Aqui eu nao olho so volume. Eu comparo tambem eficiencia. Referral nao e o maior canal em volume, mas e o melhor em conversao.

## Celula 24 - Consulta por dispositivo

Ponto principal:

- mobile tem mais volume;
- desktop tem melhores taxas de conversao.

Fala sugerida:

> Aqui aparece uma diferenca importante entre volume e eficiencia. Mobile concentra mais usuarios, mas desktop converte melhor.

## Celula 25 - Grafico de taxas por dispositivo

As barras representam taxas, nao volume absoluto. Desktop aparece melhor nas tres taxas, mas mobile segue relevante porque concentra mais visitas.

Fala sugerida:

> O grafico mostra que desktop converte melhor em todas as taxas. A oportunidade esta em mobile porque ele tem muito volume; se a conversao mobile melhorar um pouco, o impacto absoluto pode ser relevante.

## Celula 28 - Analise por pais

Leitura principal:

- BR concentra a maior parte da base;
- as taxas entre paises sao relativamente proximas;
- nao ha um pais com desvio extremo que mude toda a leitura do funil.

Fala sugerida:

> Aqui eu verifico se algum pais foge muito do padrao. O Brasil concentra o maior volume, mas as taxas entre paises ficam relativamente proximas.

## Celula 30 - Grafico mensal

Leitura principal:

- setembro apresenta a melhor taxa de compra sobre visitas;
- setembro tambem tem a melhor conversao pos-signup.

Fala sugerida:

> No grafico mensal, setembro aparece como o melhor mes em conversao para compra e conversao pos-signup. Aqui o foco e a variacao ao longo do tempo.

## Celula 33 - NPS de compradores elegiveis

A tabela mostra `respostas`, `nps_medio`, `promotores`, `passivos`, `detratores` e `nps`.

Ponto importante:

`promotores`, `passivos` e `detratores` sao contagens absolutas. O campo `nps` e o indicador percentual consolidado, em escala de -100 a 100.

Fala sugerida:

> Aqui eu trato NPS como pesquisa pos-compra. Por isso, calculo o indicador apenas para compradores elegiveis.

## Celula 34 - Grafico de NPS elegivel

Mostra o NPS dos compradores elegiveis.

Fala sugerida:

> O grafico mostra o NPS dos usuarios que compraram. Respostas associadas a usuarios sem compra nao devem ser interpretadas como NPS de nao compradores.

## Celula 37 - Conclusoes

Principais conclusoes:

- o funil tem gargalo antes do signup;
- tambem existe perda relevante depois do signup;
- referral e o canal de melhor qualidade;
- paid e social precisam de otimizacao;
- mobile e oportunidade por volume;
- NPS deve ser lido apenas para compradores elegiveis; respostas sem compra viram ponto de investigacao.

Fala sugerida:

> No final, conecto os achados as recomendacoes: escalar referral, otimizar paid, melhorar mobile, fortalecer organic, trabalhar email/lifecycle e investigar nao compradores por abandono e pesquisa qualitativa.
""".strip()


def configure_styles(document: "DocxDocument") -> None:
    for style_name in ["Normal", "List Bullet", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    document.styles["Title"].font.size = Pt(18)
    document.styles["Heading 1"].font.size = Pt(14)
    document.styles["Heading 2"].font.size = Pt(12)
    document.styles["Heading 3"].font.size = Pt(11)


def clean_inline_markdown(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def add_heading(document: "DocxDocument", text: str, level: int) -> None:
    heading = document.add_heading(clean_inline_markdown(text), level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_inline_markdown(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_bullet(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(clean_inline_markdown(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_quote(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_inline_markdown(text))
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_code(document: "DocxDocument", text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text.strip())
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def build_document() -> "DocxDocument":
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    in_code_block = False
    code_lines: list[str] = []

    for raw_line in GUIDE_MARKDOWN.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                add_code(document, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            title = document.add_heading(clean_inline_markdown(line[2:]), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = document.add_paragraph("Apoio de fala para demonstrar o notebook de EDA")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(10)
        elif line.startswith("## "):
            add_heading(document, line[3:], level=1)
        elif line.startswith("### "):
            add_heading(document, line[4:], level=2)
        elif line.startswith("> "):
            add_quote(document, line[2:])
        elif line.startswith("- "):
            add_bullet(document, line[2:])
        else:
            add_paragraph(document, line)

    if code_lines:
        add_code(document, "\n".join(code_lines))

    return document


def main() -> None:
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DOC_DIR / "guia_celulas_notebook.docx"
    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
